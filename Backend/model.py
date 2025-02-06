import pandas as pd
from docx import Document

# Load the CSV file
csv_file = 'blood_dona.csv'  # Replace with your CSV file path
data = pd.read_csv(csv_file)

# Create a Word document
document = Document()
document.add_heading('Blood Donation Data', 0)

# Add a table to the document
table = document.add_table(rows=1, cols=len(data.columns))
header_cells = table.rows[0].cells

# Add header row
for i, column_name in enumerate(data.columns):
    header_cells[i].text = column_name

# Add data rows
for row in data.itertuples(index=False):
    cells = table.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = str(value)

# Save the Word document
document.save('blood_dona.docx')
print("CSV file has been converted to DOCX.")
